from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path(r"C:\Users\SAMSUNG\Downloads\skill\tmp\docs\design-upgrade-loop\hvdc_guide_source.docx")
OUTDIR = Path(r"C:\Users\SAMSUNG\Downloads\skill\output\doc\design-upgrade\20260401-160200")
TARGET = OUTDIR / "hvdc_guide_design_upgraded.docx"

KOREAN_GUIDE = "\ud1b5\ud569 \uc6b4\uc601 \uac00\uc774\ub4dc"
KOREAN_TOC = "\ubaa9\ucc28"
KOREAN_FIGURE = "\uadf8\ub9bc "
KOREAN_DEPLOY = "\ubc30\ud3ec:"
EMOJI_WARNING = "\u26a0\ufe0f"
EMOJI_GATE = "\ud83d\udea9"
EMOJI_CHECKLIST = "\ud83d\udccb"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def ensure_run_props(run):
    if run._element.rPr is None:
        run._element.get_or_add_rPr()
    return run._element.rPr


def set_run_font(run, ascii_font: str, east_asia_font: str, size: float | None = None,
                 bold: bool | None = None, color: RGBColor | None = None, italic: bool | None = None) -> None:
    r_pr = ensure_run_props(run)
    run.font.name = ascii_font
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if italic is not None:
        run.italic = italic


def set_style_font(style, ascii_font: str, east_asia_font: str, size: float, bold: bool, color: RGBColor) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = color
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_paragraph_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = p_bdr.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        p_bdr.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")


def set_cell_margins(cell, top=70, start=90, bottom=70, end=90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def style_runs(paragraph, size: float, color: RGBColor, bold: bool | None = None, italic: bool | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, "Aptos", "Malgun Gothic", size=size, bold=bold, color=color, italic=italic)


def style_cover(paragraphs) -> None:
    for idx, para in enumerate(paragraphs):
        text = " ".join(para.text.split())
        if not text:
            continue
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if "SAMSUNG C&T / HVDC PROJECT" in text:
            para.paragraph_format.space_after = Pt(10)
            style_runs(para, 11, RGBColor(95, 107, 132), bold=True)
        elif idx in (8, 9) or ("HVDC" in text and "(TR)" in text):
            para.paragraph_format.space_after = Pt(2)
            style_runs(para, 19, RGBColor(30, 55, 95), bold=True)
        elif idx == 10 or KOREAN_GUIDE in text:
            para.paragraph_format.space_after = Pt(10)
            style_runs(para, 28, RGBColor(17, 37, 64), bold=True)
        elif "AGI Site" in text or "Mina Zayed Port" in text:
            para.paragraph_format.space_after = Pt(4)
            style_runs(para, 12, RGBColor(68, 84, 106))
        elif text.startswith("Version:") or text.startswith(KOREAN_DEPLOY):
            para.paragraph_format.space_after = Pt(3)
            style_runs(para, 9, RGBColor(110, 110, 110))


def style_document(doc: Document) -> None:
    for style_name, size, color, bold in (
        ("Title", 18, RGBColor(17, 37, 64), True),
        ("Heading 1", 16, RGBColor(17, 37, 64), True),
        ("Heading 2", 12, RGBColor(31, 78, 121), True),
        ("Heading 3", 10.5, RGBColor(47, 84, 150), True),
        ("List Paragraph", 9.5, RGBColor(45, 45, 45), False),
    ):
        try:
            set_style_font(doc.styles[style_name], "Aptos", "Malgun Gothic", size, bold, color)
        except KeyError:
            pass

    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        if section.start_type == WD_SECTION.NEW_PAGE:
            section.header_distance = Inches(0.3)
            section.footer_distance = Inches(0.35)

    toc_index = next((i for i, p in enumerate(doc.paragraphs) if p.text.strip() == KOREAN_TOC), 24)
    style_cover(doc.paragraphs[:toc_index])

    for para in doc.paragraphs:
        text = " ".join(para.text.split())
        if not text:
            continue

        style_name = para.style.name if para.style is not None else ""
        fmt = para.paragraph_format

        if style_name == "Heading 1":
            fmt.space_before = Pt(14)
            fmt.space_after = Pt(6)
            fmt.keep_with_next = True
            style_runs(para, 16, RGBColor(17, 37, 64), bold=True)
        elif style_name == "Heading 2":
            fmt.space_before = Pt(10)
            fmt.space_after = Pt(4)
            fmt.keep_with_next = True
            style_runs(para, 12, RGBColor(31, 78, 121), bold=True)
        elif style_name == "Heading 3":
            fmt.space_before = Pt(8)
            fmt.space_after = Pt(3)
            fmt.keep_with_next = True
            style_runs(para, 10.5, RGBColor(47, 84, 150), bold=True)
        elif style_name == "List Paragraph":
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(2)
            fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
            fmt.left_indent = Inches(0.18)
            fmt.first_line_indent = Inches(-0.1)
            style_runs(para, 9.5, RGBColor(45, 45, 45))
        else:
            fmt.space_before = Pt(0)
            fmt.space_after = Pt(4)
            fmt.line_spacing = 1.12
            style_runs(para, 10, RGBColor(44, 44, 44))

        if text.startswith(KOREAN_FIGURE):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fmt.space_before = Pt(3)
            fmt.space_after = Pt(6)
            style_runs(para, 9, RGBColor(89, 89, 89), italic=True)

        if text.startswith(EMOJI_WARNING):
            set_paragraph_shading(para, "FFF2CC")
            set_paragraph_border(para, "C55A11")
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(6)
            style_runs(para, 9.5, RGBColor(156, 87, 0))

        if text.startswith(EMOJI_GATE):
            set_paragraph_shading(para, "E2F0D9")
            set_paragraph_border(para, "548235")
            fmt.space_before = Pt(6)
            fmt.space_after = Pt(6)
            style_runs(para, 9.5, RGBColor(56, 87, 35), bold=True)

        if text.startswith(EMOJI_CHECKLIST):
            set_paragraph_shading(para, "DEEAF6")
            set_paragraph_border(para, "2F5597")
            fmt.space_before = Pt(8)
            fmt.space_after = Pt(4)
            style_runs(para, 10, RGBColor(31, 78, 121), bold=True)

        if text.endswith(":") and len(text) < 18:
            fmt.space_before = Pt(4)
            fmt.space_after = Pt(2)
            style_runs(para, 10, RGBColor(31, 31, 31), bold=True)

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                set_cell_margins(cell)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if r_idx == 0:
                    set_cell_shading(cell, "203864")
                elif r_idx % 2 == 0:
                    set_cell_shading(cell, "F7F9FC")

                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(2)
                    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    if r_idx == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        style_runs(para, 9, RGBColor(255, 255, 255), bold=True)
                    elif c_idx == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        style_runs(para, 9, RGBColor(45, 45, 45))
                    else:
                        style_runs(para, 9, RGBColor(45, 45, 45))


def main() -> None:
    ensure_dir(OUTDIR)
    doc = Document(str(SOURCE))
    style_document(doc)
    doc.save(str(TARGET))
    print(TARGET)


if __name__ == "__main__":
    main()
